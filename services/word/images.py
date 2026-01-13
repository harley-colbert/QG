# services/word/images.py
import os
import re
import uuid
import shutil
import zipfile
import tempfile
from typing import Dict, Any, Optional, Tuple

from PIL import Image
from lxml import etree as ET

from .constants import IMAGE_SIZE_LIMITS, logger, LOG_PREFIX

# -----------------------------------------------------------------------------
# Marker pattern: [[ Image: data.some.path ]]  (case-insensitive "Image")
# -----------------------------------------------------------------------------
IMAGE_MARKER_RE = re.compile(r"\[\[\s*[Ii]mage:?\s*([^\]]+)\s*\]\]")

# -----------------------------------------------------------------------------
# Namespaces
#   W  = WordprocessingML main
#   A  = DrawingML main
#   R  = OfficeDoc relationships (used on blip @r:embed)
#   WP = WordprocessingDrawing
#   PIC= DrawingML picture
#   PR = OPC package relationships (for word/_rels/document.xml.rels)
# -----------------------------------------------------------------------------
W   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
A   = "http://schemas.openxmlformats.org/drawingml/2006/main"
R   = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
WP  = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"
PR  = "http://schemas.openxmlformats.org/package/2006/relationships"

NSMAP = {"w": W, "a": A, "r": R, "wp": WP, "pic": PIC}

# -----------------------------------------------------------------------------
# Unit helpers
# -----------------------------------------------------------------------------
def _emu_from_mm(mm: float) -> int:
    # 1 inch = 25.4 mm; 1 inch = 914400 EMU
    return int(round(mm / 25.4 * 914400))

def _scaled_mm(image_path: str, max_w_mm: float, max_h_mm: float) -> Tuple[float, float]:
    """
    Return (w_mm, h_mm) scaled to fit within max dimensions, preserving aspect.
    Defaults to 96 dpi if missing.
    """
    with Image.open(image_path) as img:
        w_px, h_px = img.size
        dpi = (img.info.get("dpi") or (96, 96))[0] or 96
        w_mm = w_px / dpi * 25.4
        h_mm = h_px / dpi * 25.4
        scale = min(1.0, max_w_mm / w_mm if w_mm else 1.0, max_h_mm / h_mm if h_mm else 1.0)
        out_w = w_mm * scale
        out_h = h_mm * scale
        logger.debug(
            f"{LOG_PREFIX} scale: px=({w_px}x{h_px}) dpi={dpi} -> "
            f"{out_w:.2f}x{out_h:.2f}mm scale={scale:.3f}"
        )
        return out_w, out_h

# -----------------------------------------------------------------------------
# Context lookup (preserve key case; try 'data.' prefix; then lower-cased)
# -----------------------------------------------------------------------------
def _get_from_context(context: Dict[str, Any], key: str) -> Optional[str]:
    """
    Key may be 'data.a.b' or 'a.b'. We try exact first, then with 'data.' prefix,
    then lower-cased variants. Returns a string path or None.
    """
    def get_nested(ctx: Dict[str, Any], dotted: str) -> Optional[str]:
        cur = ctx
        for part in dotted.split("."):
            if isinstance(cur, dict) and part in cur:
                cur = cur[part]
            else:
                return None
        return cur if isinstance(cur, str) else None

    # exact
    val = get_nested(context, key) or get_nested(context.get("data", {}), key)
    if val:
        return val

    # try with 'data.' prefix
    if not key.startswith("data."):
        val = get_nested(context, f"data.{key}")
        if val:
            return val

    # lower variants
    low = key.lower()
    val = get_nested(context, low) or get_nested(context.get("data", {}), low)
    if val:
        return val
    if not key.startswith("data."):
        val = get_nested(context, f"data.{low}")
        if val:
            return val

    logger.warning(f"{LOG_PREFIX} image path not found for {key!r}")
    return None

# -----------------------------------------------------------------------------
# .rels helpers
# -----------------------------------------------------------------------------
def _next_rel_id(rels_root: ET.Element) -> str:
    existing = {rel.get("Id") for rel in rels_root.findall(f".//{{{PR}}}Relationship")}
    i = 1
    while True:
        rid = f"rId{i}"
        if rid not in existing:
            return rid
        i += 1

def _add_image_relationship(rels_path: str, target: str) -> str:
    """
    Append a <Relationship> in word/_rels/document.xml.rels pointing to media/<file>.
    """
    parser = ET.XMLParser(remove_blank_text=True)
    with open(rels_path, "rb") as f:
        rels_root = ET.parse(f, parser).getroot()

    rid = _next_rel_id(rels_root)
    rel_el = ET.Element(f"{{{PR}}}Relationship")
    rel_el.set("Id", rid)
    rel_el.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image")
    rel_el.set("Target", target)
    rels_root.append(rel_el)

    with open(rels_path, "wb") as f:
        ET.ElementTree(rels_root).write(
            f, xml_declaration=True, encoding="utf-8", standalone="yes"
        )
    return rid

# -----------------------------------------------------------------------------
# DrawingML picture block
# -----------------------------------------------------------------------------
def _make_drawing_inline(rId: str, cx_emu: int, cy_emu: int) -> ET.Element:
    """
    Build a minimal <wp:inline> with a <pic:pic> payload using proper namespaces.
    """
    inline = ET.Element(f"{{{WP}}}inline", nsmap=NSMAP)

    ET.SubElement(inline, f"{{{WP}}}extent", cx=str(cx_emu), cy=str(cy_emu))
    ET.SubElement(inline, f"{{{WP}}}effectExtent", l="0", t="0", r="0", b="0")
    ET.SubElement(inline, f"{{{WP}}}docPr", id="1", name=str(uuid.uuid4()))
    cnv = ET.SubElement(inline, f"{{{WP}}}cNvGraphicFramePr")
    ET.SubElement(cnv, f"{{{A}}}graphicFrameLocks", noChangeAspect="1")

    graphic = ET.SubElement(inline, f"{{{A}}}graphic")
    graphicData = ET.SubElement(graphic, f"{{{A}}}graphicData", uri=PIC)

    pic = ET.SubElement(graphicData, f"{{{PIC}}}pic")  # <-- correct, no 'pic:' prefix in the tag text
    nvPicPr = ET.SubElement(pic, f"{{{PIC}}}nvPicPr")
    ET.SubElement(nvPicPr, f"{{{PIC}}}cNvPr", id="0", name="Picture 1")
    ET.SubElement(nvPicPr, f"{{{PIC}}}cNvPicPr")

    blipFill = ET.SubElement(pic, f"{{{PIC}}}blipFill")
    blip = ET.SubElement(blipFill, f"{{{A}}}blip")
    blip.set(f"{{{R}}}embed", rId)
    stretch = ET.SubElement(blipFill, f"{{{A}}}stretch")
    ET.SubElement(stretch, f"{{{A}}}fillRect")

    spPr = ET.SubElement(pic, f"{{{PIC}}}spPr")
    xfrm = ET.SubElement(spPr, f"{{{A}}}xfrm")
    ET.SubElement(xfrm, f"{{{A}}}off", x="0", y="0")
    ET.SubElement(xfrm, f"{{{A}}}ext", cx=str(cx_emu), cy=str(cy_emu))
    prstGeom = ET.SubElement(spPr, f"{{{A}}}prstGeom", prst="rect")
    ET.SubElement(prstGeom, f"{{{A}}}avLst")

    return inline

def _replace_marker_para_with_image(p: ET.Element, rid: str, cx_emu: int, cy_emu: int) -> None:
    """
    Replace the entire paragraph content with a single run containing the drawing.
    """
    for child in list(p):
        p.remove(child)
    r = ET.SubElement(p, f"{{{W}}}r")
    drawing = ET.SubElement(r, f"{{{W}}}drawing")
    # NOTE: There was a stray '}' in some drafts; ensure correct tag:
    # drawing = ET.SubElement(r, f"{{{W}}}drawing")
    # Fix immediately:
    p.remove(r)  # remove the bad one
    r = ET.SubElement(p, f"{{{W}}}r")
    drawing = ET.SubElement(r, f"{{{W}}}drawing")
    drawing.append(_make_drawing_inline(rid, cx_emu, cy_emu))

# (helper) iterate paragraphs and collect all text content
def _iter_paragraphs(doc_root: ET.Element):
    for p in doc_root.findall(".//w:p", namespaces=NSMAP):
        yield p

def _para_text(p: ET.Element) -> str:
    texts = p.findall(".//w:t", namespaces=NSMAP)
    return "".join(t.text or "" for t in texts)

# -----------------------------------------------------------------------------
# Public: XML-based replacement (recommended)
# -----------------------------------------------------------------------------
def replace_image_markers_xml(docx_path: str, context: Dict[str, Any]) -> None:
    """
    Open the rendered .docx, replace [[ Image: ... ]] markers with inline images by
    editing document.xml and document.xml.rels directly. This avoids python-docx and
    preserves any Subdoc content docxtpl already injected.
    """
    # 0) temp working dir + copy input (some environments prefer not touching the file in place)
    workdir = tempfile.mkdtemp(prefix="imgxml_")
    tmp_docx = os.path.join(workdir, "in.docx")
    shutil.copy2(docx_path, tmp_docx)

    unzip_dir = os.path.join(workdir, "unzipped")
    os.makedirs(unzip_dir, exist_ok=True)

    with zipfile.ZipFile(tmp_docx, "r") as zf:
        zf.extractall(unzip_dir)

    doc_xml = os.path.join(unzip_dir, "word", "document.xml")
    rels_xml = os.path.join(unzip_dir, "word", "_rels", "document.xml.rels")
    media_dir = os.path.join(unzip_dir, "word", "media")
    os.makedirs(media_dir, exist_ok=True)

    parser = ET.XMLParser(remove_blank_text=True)
    with open(doc_xml, "rb") as f:
        doc_root = ET.parse(f, parser).getroot()

    changed = False

    for p in list(_iter_paragraphs(doc_root)):
        text = _para_text(p)
        m = IMAGE_MARKER_RE.search(text or "")
        if not m:
            continue

        logger.debug(f"{LOG_PREFIX} images(xml): para matches={[m.group(1)]}")
        orig_key = m.group(1).strip()
        image_key = orig_key if orig_key.startswith("data.") else f"data.{orig_key}"
        image_path = _get_from_context(context, image_key)
        if not image_path or not os.path.isfile(image_path):
            logger.error(f"{LOG_PREFIX} missing image path for {image_key}")
            continue

        # Copy the image into /word/media/ and add a relationship
        ext = os.path.splitext(image_path)[1].lower() or ".png"
        target_name = f"{uuid.uuid4().hex}{ext}"
        target_rel = f"media/{target_name}"
        shutil.copy2(image_path, os.path.join(media_dir, target_name))

        rid = _add_image_relationship(rels_xml, target_rel)

        # Size in EMUs
        short_key = image_key.replace("data.", "")
        if short_key in IMAGE_SIZE_LIMITS:
            max_w, max_h = IMAGE_SIZE_LIMITS[short_key]
            w_mm, h_mm = _scaled_mm(image_path, max_w, max_h)
        else:
            # "natural" size @ 96dpi
            w_mm, h_mm = _scaled_mm(image_path, 99999, 99999)

        cx, cy = _emu_from_mm(w_mm), _emu_from_mm(h_mm)

        # Replace the whole paragraph with the picture drawing
        _replace_marker_para_with_image(p, rid, cx, cy)
        logger.info(f"{LOG_PREFIX} inserted image (xml) key={image_key} size={w_mm:.1f}x{h_mm:.1f}mm")
        changed = True

    if changed:
        with open(doc_xml, "wb") as f:
            ET.ElementTree(doc_root).write(f, xml_declaration=True, encoding="utf-8", standalone="yes")

        # Rezip package back to the original path
        tmp_out = os.path.join(workdir, "out.docx")
        with zipfile.ZipFile(tmp_out, "w", zipfile.ZIP_DEFLATED) as z:
            for root, _, files in os.walk(unzip_dir):
                for name in files:
                    full = os.path.join(root, name)
                    arcname = os.path.relpath(full, unzip_dir)
                    z.write(full, arcname)

        shutil.copy2(tmp_out, docx_path)

    shutil.rmtree(workdir, ignore_errors=True)

# -----------------------------------------------------------------------------
# Legacy python-docx-based replacement (kept for backward-compatibility).
# NOTE: This mutates runs and can conflict with Subdocs. Prefer XML path above.
# -----------------------------------------------------------------------------
try:
    from docx import Document as _PyDocxDocument
    from docx.shared import Mm as _PyDocxMm
except Exception:  # pragma: no cover
    _PyDocxDocument = None
    _PyDocxMm = None

def replace_image_markers(doc, context: Dict[str, Any]) -> None:
    """
    Backwards compatible shim. If 'doc' is a python-docx Document, do the old run-based
    replacement. Otherwise, if it's a str (path), fall back to XML method.
    """
    if isinstance(doc, str):
        # Treat as a path to a docx; use XML method to keep Subdocs safe.
        replace_image_markers_xml(doc, context)
        return

    if _PyDocxDocument is None or not hasattr(doc, "paragraphs"):
        logger.warning(f"{LOG_PREFIX} replace_image_markers fallback: unsupported doc type; using XML method if possible.")
        return

    # NB: This path is *not* recommended when you’re inserting Subdocs.
    for paragraph in doc.paragraphs:
        matches = IMAGE_MARKER_RE.findall(paragraph.text or "")
        if not matches:
            continue
        logger.debug(f"{LOG_PREFIX} images: para matches={matches}")
        for match in matches:
            orig_key = match.strip()
            image_key = orig_key if orig_key.startswith("data.") else f"data.{orig_key}"
            image_path = _get_from_context(context, image_key)
            if not image_path or not os.path.isfile(image_path):
                logger.error(f"{LOG_PREFIX} missing image path for {image_key}")
                continue

            marker_text = f"[[ Image: {orig_key} ]]"
            paragraph.text = (paragraph.text or "").replace(marker_text, "")

            run = paragraph.add_run()
            short_key = image_key.replace("data.", "")
            if short_key in IMAGE_SIZE_LIMITS and _PyDocxMm is not None:
                max_w, max_h = IMAGE_SIZE_LIMITS[short_key]
                width_mm, height_mm = _scaled_mm(image_path, max_w, max_h)
                run.add_picture(image_path, width=_PyDocxMm(width_mm), height=_PyDocxMm(height_mm))
                logger.info(f"{LOG_PREFIX} inserted image key={image_key} size={width_mm:.1f}x{height_mm:.1f}mm")
            else:
                run.add_picture(image_path)
                logger.info(f"{LOG_PREFIX} inserted image key={image_key} default size")
