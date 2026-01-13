# services/word/html_shim.py
# HTML -> Word bridge using html4docx; accepts either python-docx Document or docxtpl.Subdoc.

from __future__ import annotations
from typing import Optional, Iterable, Any

from docx.document import Document as _DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.text.paragraph import Paragraph

from .constants import logger, LOG_PREFIX

try:
    import html4docx  # type: ignore
except Exception as e:  # pragma: no cover
    raise ImportError(
        f"{LOG_PREFIX} html-for-docx is required. Install with:\n"
        "    pip install html-for-docx\n"
        f"Original import error: {e}"
    )

_parser = html4docx.HtmlToDocx()

# ---------- toggles ----------
FORCE_LIST_PARAGRAPH_STYLE = False  # keep False unless you really want to flatten
SYNC_STYLE_TO_LEVEL = True          # recommended
RELATIVE_TO_ANCHOR_LEVEL = True     # style sync only (numbering already set by html4docx)


def ensure_html_wrapper_initialized() -> None:
    logger.debug(f"{LOG_PREFIX} html_shim: initialized")


def add_html_wrapper(doc_or_subdoc: Any, html_in: str) -> None:
    """
    Convert HTML to Word. Accepts either:
      - python-docx Document
      - docxtpl.Subdoc  (we unwrap to its .docx Document)
    Preserves nested lists and (optionally) syncs style names to levels.
    """
    logger.debug(f"{LOG_PREFIX} html_shim.add_html_wrapper: start")

    # Coerce to python-docx Document (what html4docx requires)
    doc = _coerce_to_document(doc_or_subdoc)

    _ensure_list_styles(doc)

    before_ids = {id(p) for p in _iter_all_paragraphs(doc)}
    _parser.add_html_to_document(html_in, doc)
    logger.debug(f"{LOG_PREFIX} html-for-docx conversion complete")

    after_paras = list(_iter_all_paragraphs(doc))
    inserted_ids = {id(p) for p in after_paras if id(p) not in before_ids}

    anchor = _infer_anchor_from_doc(before_ids, after_paras) if RELATIVE_TO_ANCHOR_LEVEL else None

    if FORCE_LIST_PARAGRAPH_STYLE:
        _relabel_list_styles_to_list_paragraph(doc, inserted_ids)

    if SYNC_STYLE_TO_LEVEL:
        _sync_style_name_with_level(doc, inserted_ids, anchor)

    logger.debug(f"{LOG_PREFIX} html_shim.add_html_wrapper: done")


# ---------- helpers ----------
def _coerce_to_document(obj: Any) -> _DocxDocument:
    """
    Accept python-docx Document OR docxtpl.Subdoc. Return a python-docx Document.
    """
    if isinstance(obj, _DocxDocument):
        return obj

    # Lazy import to avoid hard dependency if you never pass Subdoc
    try:
        from docxtpl.subdoc import Subdoc  # type: ignore
    except Exception:
        Subdoc = None  # type: ignore

    # docxtpl.Subdoc has a .docx attribute (python-docx Document)
    if Subdoc is not None and isinstance(obj, Subdoc):
        if hasattr(obj, "docx") and isinstance(obj.docx, _DocxDocument):
            return obj.docx  # <-- critical: unwrap Subdoc to Document
        raise TypeError(f"{LOG_PREFIX} Subdoc detected but missing a valid .docx Document")

    # Some wrappers might carry a .docx as well—gracefully support that.
    if hasattr(obj, "docx") and isinstance(getattr(obj, "docx"), _DocxDocument):
        return getattr(obj, "docx")

    raise TypeError(
        f"{LOG_PREFIX} add_html_wrapper needs a python-docx Document or docxtpl.Subdoc; "
        f"got {type(obj)!r}"
    )


def _iter_all_paragraphs(doc: _DocxDocument) -> Iterable[Paragraph]:
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _ensure_list_styles(doc: _DocxDocument) -> None:
    styles = doc.styles
    for base, count in (("List Bullet", 9), ("List Number", 9)):
        for i in range(1, count + 1):
            
            if i<2:
                name = f"{base}"
            else:
                name = f"{base} {i}"
            if name not in styles:
                s = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
                s.quick_style = True
                logger.debug(f"{LOG_PREFIX} created missing style: {name}")


def _get_num_ilvl(paragraph: Paragraph) -> Optional[int]:
    pPr = paragraph._p.pPr
    if pPr is None or pPr.numPr is None:
        return None
    ilvl = pPr.numPr.ilvl
    if ilvl is None or ilvl.val is None:
        return None
    try:
        return int(ilvl.val)
    except Exception:
        return None


def _is_numbered(paragraph: Paragraph) -> bool:
    pPr = paragraph._p.pPr
    return bool(pPr is not None and pPr.numPr is not None)


def _paragraph_kind(paragraph: Paragraph) -> str:
    """
    Heuristic: detect 'bullet' vs 'number'.
    We primarily use style name as a hint after html4docx mapping.
    """
    style = paragraph.style.name if paragraph.style else ""
    if "Number" in style:
        return "number"
    if "Bullet" in style:
        return "bullet"
    # Fallback: default to bullet to avoid wrong Number styling on bullets
    return "bullet"


def _sync_style_name_with_level(doc: _DocxDocument,
                                inserted_ids: set[int],
                                anchor: Optional[Paragraph]) -> None:
    anchor_level = _get_num_ilvl(anchor) if anchor is not None else None

    for p in _iter_all_paragraphs(doc):
        if id(p) not in inserted_ids:
            continue
        if not _is_numbered(p):
            continue

        ilvl = _get_num_ilvl(p)
        if ilvl is None:
            continue

        effective_level = ilvl
        if anchor_level is not None:
            # numbering ilvl already correct; we only reflect it in style text
            effective_level = ilvl

        style_index = max(1, min(9, effective_level + 1))  # 0-based ilvl -> 1..9
        kind = _paragraph_kind(p)
        
        
        if style_index <2:
            style_name = f"List Number 1" if kind == "number" else f"List Bullet"
        else:
            style_name = f"List Number {style_index}" if kind == "number" else f"List Bullet {style_index}"
        _set_paragraph_style_name(p, style_name)
        logger.debug(f"{LOG_PREFIX} synced style: '{p.text[:40]}' -> {style_name} (ilvl={ilvl}, kind={kind})")


def _set_paragraph_style_name(paragraph: Paragraph, style_name: str) -> None:
    try:
        paragraph.style = style_name
    except KeyError:
        s = paragraph.part.document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        s.quick_style = True
        paragraph.style = style_name


def _relabel_list_styles_to_list_paragraph(doc: _DocxDocument, inserted_ids: set[int]) -> None:
    for p in _iter_all_paragraphs(doc):
        if id(p) not in inserted_ids:
            continue
        if _is_numbered(p):
            _set_paragraph_style_name(p, "List Paragraph")


def _infer_anchor_from_doc(before_ids: set[int], after_paras: list[Paragraph]) -> Optional[Paragraph]:
    first_new_i = None
    for i, p in enumerate(after_paras):
        if id(p) not in before_ids:
            first_new_i = i
            break
    if first_new_i is None:
        return None
    return after_paras[first_new_i - 1] if first_new_i > 0 else None
