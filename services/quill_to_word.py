import os
import re
import logging
import time
import tempfile
from typing import Any, Callable, Optional, Dict, List, Union

import pythoncom
import win32com.client as win32
from win32com.client import constants as c

from docxtpl import DocxTemplate  # Subdoc is created from docxtpl instance via new_subdoc()
from docx import Document
from docx.shared import Mm
from bs4 import BeautifulSoup
from docx.enum.text import WD_BREAK

from html2docx import add_html_to_document  # <- key piece for HTML→Word lists/paras

WD_WITHIN_TABLE = 12  # WdInformation.wdWithInTable
LOG_PREFIX = "[QUILL2DOC]"

# ---------------------------------------------------------------------------
# Sizing rules for images keyed by your content keys (mm)
# ---------------------------------------------------------------------------
IMAGE_SIZE_LIMITS = {
    "customercontact.logo": (127, 45.72),
    "titleImage": (95.25, 57.15),
    "systemLayout.elevation": (158.75, 152.4),
    "systemLayout.end": (158.75, 152.4),
    "systemLayout.iso": (158.75, 152.4),
    "systemLayout.top": (158.75, 152.4),
    "systemLayout.title": (95.25, 57.15),
}

# ---------------------------------------------------------------------------
# Only these exact dotted paths are allowed to stay "rich" (HTML/Delta).
# Everything else will be sanitized to plain text.
# ---------------------------------------------------------------------------
EXCEPTION_HTML_FIELDS = {
    "data.zonefunctions.guardingdescription",
    "data.riskdescription.1",
    "data.zonefunctions.controlsdescription",
}

logger = logging.getLogger("quill2doc")
logger.setLevel(logging.DEBUG)
if not logger.hasHandlers():
    handler = logging.StreamHandler()
    # force our identifier on every line, plus time/module/level
    formatter = logging.Formatter(
        f"%(asctime)s - {LOG_PREFIX} - %(name)s - %(levelname)s - %(message)s"
    )
    handler.setFormatter(formatter)
    logger.addHandler(handler)


# -----------------------------------------------------------------------------
# Utilities
# -----------------------------------------------------------------------------
def strip_markup(text: Any) -> Any:
    """Strip all HTML/XML tags from a string."""
    if not isinstance(text, str):
        return text
    stripped = BeautifulSoup(text, "html.parser").get_text()
    logger.debug(f"{LOG_PREFIX} strip_markup: input_len={len(text)} output_len={len(stripped)}")
    return stripped


def get_scaled_dimensions(image_path: str, max_width_mm: float, max_height_mm: float):
    from PIL import Image
    logger.debug(f"{LOG_PREFIX} get_scaled_dimensions: path={image_path!r} max=({max_width_mm},{max_height_mm})mm")
    with Image.open(image_path) as img:
        img_width, img_height = img.size
        dpi = img.info.get("dpi", (96, 96))[0]
        width_mm = img_width / dpi * 25.4
        height_mm = img_height / dpi * 25.4
        width_scale = max_width_mm / width_mm
        height_scale = max_height_mm / height_mm
        scale = min(1, width_scale, height_scale)
        out_w, out_h = width_mm * scale, height_mm * scale
        logger.debug(f"{LOG_PREFIX} get_scaled_dimensions: img=({img_width}x{img_height})px @ {dpi}dpi -> ({out_w:.2f} x {out_h:.2f})mm scale={scale:.3f}")
        return out_w, out_h


# -----------------------------------------------------------------------------
# Optional: Quill Delta → HTML (handles paragraphs, ordered/bulleted lists,
# and common inline styles/links). If your inputs are already HTML, this is
# just unused—but it lets you post raw Delta too.
# -----------------------------------------------------------------------------
def _looks_like_quill_delta(s: str) -> bool:
    test = s.strip()
    looks = test.startswith("{") and '"ops"' in test
    logger.debug(f"{LOG_PREFIX} _looks_like_quill_delta: looks={looks} sample={test[:80].replace(chr(10),' ')}...")
    return looks


def delta_to_html(delta: Union[str, Dict[str, Any]]) -> str:
    import json

    logger.debug(f"{LOG_PREFIX} delta_to_html: begin type={type(delta).__name__}")
    if isinstance(delta, str):
        try:
            raw_preview = delta[:120].replace("\n", " ")
            logger.debug(f"{LOG_PREFIX} delta_to_html: str input preview={raw_preview!r}")
            delta = json.loads(delta)
            logger.debug(f"{LOG_PREFIX} delta_to_html: parsed JSON ok")
        except Exception as e:
            logger.warning(f"{LOG_PREFIX} delta_to_html: invalid JSON; fallback to escaped paragraph. err={e}")
            return f"<p>{strip_markup(delta)}</p>"

    ops: List[Dict[str, Any]] = delta.get("ops", [])
    logger.debug(f"{LOG_PREFIX} delta_to_html: ops_count={len(ops)}")
    # Split into lines by '\n' (Quill block separator)
    lines: List[Dict[str, Any]] = []
    buffer_text = ""
    buffer_attrs: Dict[str, Any] = {}

    def flush_line(attrs: Dict[str, Any]):
        nonlocal buffer_text
        text = buffer_text
        buffer_text = ""
        if text.endswith("\n"):
            text = text[:-1]
        lines.append({"text": text, "attrs": attrs.copy()})

    for idx, op in enumerate(ops):
        insert = op.get("insert", "")
        attributes = op.get("attributes", {}) or {}
        if isinstance(insert, dict):
            # Simple placeholder for unsupported embeds (extend if needed for images)
            if "image" in insert:
                buffer_text += f"[image: {insert['image']}]"
                logger.debug(f"{LOG_PREFIX} delta_to_html: op#{idx} embed image")
            else:
                buffer_text += "[embed]"
                logger.debug(f"{LOG_PREFIX} delta_to_html: op#{idx} generic embed")
            continue
        # Text may include one or more '\n'
        parts = str(insert).split("\n")
        for i, part in enumerate(parts):
            if i > 0:
                flush_line(buffer_attrs)
            buffer_attrs = attributes.copy()
            buffer_text += part

    if buffer_text:
        flush_line(buffer_attrs)

    logger.debug(f"{LOG_PREFIX} delta_to_html: lines_count={len(lines)}")

    # Build HTML with correct list grouping
    html_parts: List[str] = []
    open_list_type: Optional[str] = None  # "ul" | "ol" | None
    li_count = 0
    p_count = 0

    def _inline_spans(text: str, attrs: Dict[str, Any]) -> str:
        esc = (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )
        if attrs.get("link"):
            esc = f'<a href="{attrs["link"]}">{esc}</a>'
        if attrs.get("bold"):
            esc = f"<strong>{esc}</strong>"
        if attrs.get("italic"):
            esc = f"<em>{esc}</em>"
        if attrs.get("underline"):
            esc = f"<u>{esc}</u>"
        if attrs.get("strike"):
            esc = f"<s>{esc}</s>"
        return esc

    def _close_list():
        nonlocal open_list_type
        if open_list_type:
            html_parts.append(f"</{open_list_type}>")
            logger.debug(f"{LOG_PREFIX} delta_to_html: close_list type={open_list_type}")
            open_list_type = None

    for i, line in enumerate(lines):
        text = line["text"]
        attrs = line["attrs"]
        lst = attrs.get("list")
        if lst in ("ordered", "bullet"):
            desired = "ol" if lst == "ordered" else "ul"
            if open_list_type != desired:
                _close_list()
                open_list_type = desired
                html_parts.append(f"<{desired}>")
                logger.debug(f"{LOG_PREFIX} delta_to_html: open_list type={desired}")
            html_parts.append(f"<li>{_inline_spans(text, attrs)}</li>")
            li_count += 1
        else:
            if open_list_type:
                _close_list()
            content = _inline_spans(text if text.strip() else "\u00A0", attrs)
            html_parts.append(f"<p>{content}</p>")
            p_count += 1

    if open_list_type:
        _close_list()

    html = "\n".join(html_parts)
    logger.debug(f"{LOG_PREFIX} delta_to_html: done p_count={p_count} li_count={li_count} html_len={len(html)}")
    return html


# -----------------------------------------------------------------------------
# Sanitization: everything becomes plain text EXCEPT exact paths in
# EXCEPTION_HTML_FIELDS (those we leave as-is to become Subdocs).
# -----------------------------------------------------------------------------
def sanitize_except_exceptions(obj: Any, parent_path: Optional[List[str]] = None) -> Any:
    if parent_path is None:
        parent_path = []
    if isinstance(obj, dict):
        clean = {}
        for k, v in obj.items():
            path = parent_path + [k]
            full_path = ".".join(path).lower()
            if full_path in EXCEPTION_HTML_FIELDS:
                logger.debug(f"{LOG_PREFIX} sanitize: KEEP-RICH path={full_path}")
                clean[k] = v  # leave HTML/Delta as-is
            else:
                clean[k] = sanitize_except_exceptions(v, path)
        return clean
    elif isinstance(obj, list):
        return [sanitize_except_exceptions(item, parent_path) for item in obj]
    elif isinstance(obj, str):
        s = strip_markup(obj)  # convert to plain
        if s != obj:
            logger.debug(f"{LOG_PREFIX} sanitize: STRIPPED len_in={len(obj)} len_out={len(s)}")
        return s
    else:
        return obj


# -----------------------------------------------------------------------------
# Main service
# -----------------------------------------------------------------------------
class WordSubmissionHandler:
    def __init__(
        self,
        template_path: str,
        status_callback: Optional[Callable[[str], None]] = None,
    ) -> None:
        self.template_path = template_path
        self.status_callback = status_callback
        self.doc: Optional[Document] = None
        self.tpl: Optional[DocxTemplate] = None
        logger.debug(f"{LOG_PREFIX} __init__: template_path={template_path!r}")

    # ---------- misc ----------
    def send_status(self, message: str) -> None:
        msg = f"{LOG_PREFIX} {message}"
        if self.status_callback:
            self.status_callback(msg)
        logger.info(msg)

    def load_template(self) -> None:
        self.send_status("Loading Word template...")
        if not os.path.isfile(self.template_path):
            logger.error(f"{LOG_PREFIX} load_template: not found {self.template_path!r}")
            raise FileNotFoundError(f"Template file not found: {self.template_path}")
        try:
            self.tpl = DocxTemplate(self.template_path)
            self.send_status("Template loaded successfully.")
        except Exception as e:
            logger.exception(f"{LOG_PREFIX} Error loading template.")
            raise e

    # ---------- core: transform rich fields into Subdocs ----------
    def _build_subdoc_from_html(self, html: str):
        """
        Create a docxtpl Subdoc and inject HTML into it using html2docx.
        """
        if not self.tpl:
            raise RuntimeError("Template not loaded.")
        logger.debug(f"{LOG_PREFIX} _build_subdoc_from_html: html_len={len(html)} preview={html[:120]!r}")
        sub = self.tpl.new_subdoc()
        add_html_to_document(sub, html)
        logger.debug(f"{LOG_PREFIX} _build_subdoc_from_html: subdoc OK")
        return sub

    def _transform_exceptions_to_subdocs(
        self, obj: Any, parent_path: Optional[List[str]] = None
    ) -> Any:
        """
        Walk the sanitized context and, for exact dotted paths in EXCEPTION_HTML_FIELDS,
        convert the string (HTML or Quill Delta) into a Subdoc object.
        """
        if parent_path is None:
            parent_path = []

        if isinstance(obj, dict):
            out = {}
            for k, v in obj.items():
                path = parent_path + [k]
                full = ".".join(path).lower()
                if full in EXCEPTION_HTML_FIELDS and isinstance(v, str):
                    logger.debug(f"{LOG_PREFIX} _transform: RICH path={full} val_len={len(v)} looksDelta={_looks_like_quill_delta(v)}")
                    html = v
                    if _looks_like_quill_delta(v):
                        try:
                            html = delta_to_html(v)
                        except Exception as e:
                            logger.warning(f"{LOG_PREFIX} Delta→HTML failed at {full}: {e}; falling back to plain text.")
                            html = f"<p>{strip_markup(v)}</p>"
                    out[k] = self._build_subdoc_from_html(html)
                else:
                    out[k] = self._transform_exceptions_to_subdocs(v, path)
            return out
        elif isinstance(obj, list):
            return [self._transform_exceptions_to_subdocs(v, parent_path) for v in obj]
        else:
            return obj

    # ---------- public API ----------
    def render_document(self, context: dict, output_path: str):
        logger.info(f"{LOG_PREFIX} [Render_Content] keys={list(context.keys())} output_path={output_path!r}")

        # 1) Sanitize (only exceptions retain "rich" values)
        self.send_status("Step 1: Sanitizing context (non-exception fields → plain)...")
        sanitized_context = sanitize_except_exceptions(context)
        #logger.debug(f"{LOG_PREFIX} [Sanitized_Content] snapshot={str(sanitized_context)[:350]}...")

        # 2) Load template
        self.send_status("Step 2: Loading DOCX template...")
        self.load_template()

        # 3) Convert exception fields into Subdocs (HTML/Delta → real Word content)
        self.send_status("Step 3: Converting rich fields to Subdocs (HTML → Word)...")
        context_for_tpl = self._transform_exceptions_to_subdocs(sanitized_context)
        logger.debug(f"{LOG_PREFIX} context_for_tpl ready (contains Subdocs)")

        # 4) Render once with docxtpl (Subdocs land exactly where placeholders are)
        self.send_status("Step 4: Rendering DOCX template with docxtpl...")
        self.tpl.render(context_for_tpl)

        temp_doc_path = os.path.join(tempfile.gettempdir(), "temp_rendered.docx")
        self.send_status(f"Step 5: Saving rendered document to {temp_doc_path}")
        self.tpl.save(temp_doc_path)

        # 5) Replace image markers in-place
        self.send_status("Step 6: Replacing image markers...")
        self.doc = Document(temp_doc_path)
        self.replace_image_markers(context)

        temp_after_images = os.path.join(tempfile.gettempdir(), "temp_after_images.docx")
        self.send_status(f"Step 7: Saving after image replacement to {temp_after_images}")
        self.doc.save(temp_after_images)

        # 6) Insert page breaks (COM-based pagination pass)
        self.send_status("Step 8: Inserting page breaks based on vertical position...")
        #self.insert_page_breaks_by_vertical_position(temp_after_images, 7.25)

        self.doc = Document(temp_after_images)
        self.send_status(f"Step 9: Final save to {output_path}...")
        self.doc.save(output_path)
        self.send_status(f"Document fully processed and saved to {output_path}")

    # ---------- image replacement ----------
    def get_image_path_from_context(self, context: dict, key: str) -> Optional[str]:
        logger.debug(f"{LOG_PREFIX} [get_image_path_from_context] START for key={repr(key)}")

        def get_nested(ctx, k):
            logger.debug(f"{LOG_PREFIX}   [get_nested] type={type(ctx)} key={k}")
            keys = k.split(".")
            curr = ctx
            try:
                for p in keys:
                    if isinstance(curr, dict) and p in curr:
                        curr = curr[p]
                        logger.debug(f"{LOG_PREFIX}     [get_nested] found {p}: type={type(curr)}")
                    else:
                        logger.debug(f"{LOG_PREFIX}     [get_nested] {p} not found or curr not dict")
                        return None
                if isinstance(curr, str):
                    logger.debug(f"{LOG_PREFIX}   [get_nested] final string: {curr!r}")
                    return curr
                else:
                    logger.debug(f"{LOG_PREFIX}   [get_nested] final non-string: {type(curr)}")
                    return None
            except Exception as e:
                logger.debug(f"{LOG_PREFIX}   [get_nested] exception: {e}")
                return None

        logger.debug(f"{LOG_PREFIX}   try nested in context")
        val = get_nested(context, key)
        if val:
            logger.info(f"{LOG_PREFIX}   found image in context: {val!r}")
            return val

        logger.debug(f"{LOG_PREFIX}   try nested in context['data']")
        if "data" in context and isinstance(context["data"], dict):
            val = get_nested(context["data"], key)
            if val:
                logger.info(f"{LOG_PREFIX}   found image in context['data']: {val!r}")
                return val

        logger.debug(f"{LOG_PREFIX}   try flat key in context")
        if key in context and isinstance(context[key], str):
            logger.info(f"{LOG_PREFIX}   found flat in context: {context[key]!r}")
            return context[key]

        logger.debug(f"{LOG_PREFIX}   try flat key in context['data']")
        if "data" in context and key in context["data"] and isinstance(
            context["data"][key], str
        ):
            logger.info(f"{LOG_PREFIX}   found flat in context['data']: {context['data'][key]!r}")
            return context["data"][key]

        logger.debug(f"{LOG_PREFIX}   try lowercase key in context")
        if key.lower() in context and isinstance(context[key.lower()], str):
            logger.info(f"{LOG_PREFIX}   found lowercase in context: {context[key.lower()]!r}")
            return context[key.lower()]

        logger.debug(f"{LOG_PREFIX}   try lowercase key in context['data']")
        if "data" in context and key.lower() in context["data"] and isinstance(
            context["data"][key.lower()], str
        ):
            logger.info(f"{LOG_PREFIX}   found lowercase in context['data']: {context['data'][key.lower()]!r}")
            return context["data"][key.lower()]

        logger.warning(f"{LOG_PREFIX} [get_image_path_from_context] not found for key: {key!r}")
        return None

    def replace_image_markers(self, context: dict):
        """
        Replaces all [[ Image: key ]] markers in the document with the specified images.
        """
        self.send_status("Replacing image markers...")
        image_marker_pattern = re.compile(r"\[\[\s*[Ii]mage:?\s*([^\]]+)\s*\]\]")
        try:
            if self.doc is None:
                raise ValueError("Document not loaded. Call load_template() first.")

            para_count = 0
            marker_total = 0
            for paragraph in self.doc.paragraphs:
                para_count += 1
                matches = image_marker_pattern.findall(paragraph.text)
                if matches:
                    logger.debug(f"{LOG_PREFIX} replace_image_markers: para#{para_count} matches={matches}")
                for match in matches:
                    marker_total += 1
                    orig_key = match.strip()
                    image_key = (
                        orig_key if orig_key.startswith("data.") else f"data.{orig_key}"
                    )
                    logger.debug(
                        f"{LOG_PREFIX} replace_image_markers: lookup image_key={image_key} (orig={orig_key})"
                    )
                    image_path = self.get_image_path_from_context(context, image_key)
                    logger.debug(
                        f"{LOG_PREFIX} replace_image_markers: resolved path -> {image_path!r}"
                    )
                    if not image_path or not os.path.isfile(image_path):
                        logger.error(f"{LOG_PREFIX} replace_image_markers: missing path for {image_key}")
                        continue

                    marker_text = f"[[ Image: {orig_key} ]]"
                    paragraph.text = paragraph.text.replace(marker_text, "")

                    run = paragraph.add_run()
                    short_key = image_key.replace("data.", "")
                    if short_key in IMAGE_SIZE_LIMITS:
                        max_w, max_h = IMAGE_SIZE_LIMITS[short_key]
                        width_mm, height_mm = get_scaled_dimensions(
                            image_path, max_w, max_h
                        )
                        run.add_picture(
                            image_path, width=Mm(width_mm), height=Mm(height_mm)
                        )
                        logger.info(
                            f"{LOG_PREFIX} insert image: key='{image_key}' scaled={width_mm:.1f}x{height_mm:.1f}mm"
                        )
                    else:
                        run.add_picture(image_path)
                        logger.info(
                            f"{LOG_PREFIX} insert image: key='{image_key}' default size"
                        )

            logger.info(f"{LOG_PREFIX} replace_image_markers: finished markers={marker_total}")
            self.send_status("Image markers replaced successfully.")
        except Exception as e:
            logger.exception(f"{LOG_PREFIX} Error replacing image markers.")
            raise e

    # ---------- page break pass ----------
    def insert_page_breaks_by_vertical_position(self, doc_path, y_in_inches=7.0):
        from pywintypes import com_error

        c = win32.constants
        threshold_points = y_in_inches * 72
        max_page = 14

        logger.debug(f"{LOG_PREFIX} insert_page_breaks: open Word doc={doc_path!r} y_in={y_in_inches}in threshold={threshold_points}pt")
        word = win32.Dispatch("Word.Application")
        word.Visible = False

        pass_count = 0
        inserted_break = True

        while inserted_break:
            inserted_break = False
            try:
                doc = word.Documents.Open(doc_path)
                doc.Repaginate()
                count = doc.Paragraphs.Count
                print(f"\n===== {LOG_PREFIX} Pass #{pass_count + 1} | Paragraph count: {count} =====")
                i = 0
                while i < count - 1:
                    para = doc.Paragraphs(i + 1)
                    rng = para.Range
                    para_text = (
                        rng.Text.replace("\r", "\\r").replace("\n", "\\n").strip()
                    )
                    style_name = str(para.Style)
                    try:
                        y = rng.Information(c.wdVerticalPositionRelativeToPage)
                        page_num = rng.Information(c.wdActiveEndPageNumber)
                    except Exception:
                        y = None
                        page_num = None

                    next_para = doc.Paragraphs(i + 2)
                    next_rng = next_para.Range
                    next_para_text = (
                        next_rng.Text.replace("\r", "\\r")
                        .replace("\n", "\\n")
                        .strip()
                    )
                    try:
                        next_y = next_rng.Information(c.wdVerticalPositionRelativeToPage)
                    except Exception:
                        next_y = None

                    is_heading_1 = style_name.lower() == "heading 1"
                    is_bold_start = False
                    bold_flag = None
                    try:
                        first_word_rng = rng.Words(1)
                        bold_flag = first_word_rng.Font.Bold
                        is_bold_start = bool(bold_flag)
                    except Exception:
                        pass

                    print(
                        f"\n[{i+1:>3}] {LOG_PREFIX} '{para_text[:40]}' | Style: '{style_name}' | "
                        f"Page={page_num} | y={repr(y)} | Next_y={repr(next_y)} | "
                        f"Heading1={is_heading_1} | Bold={bold_flag}"
                    )

                    if rng.Information(WD_WITHIN_TABLE):
                        print(f"  > {LOG_PREFIX} In table. Skipping.")
                        i += 1
                        continue

                    if page_num is not None and page_num > max_page:
                        print(f"[{i+1:>3}] {LOG_PREFIX} Page {page_num} > {max_page}. Skipping rest.")
                        break

                    if is_heading_1 and y is not None and y > threshold_points:
                        print(
                            f"  > {LOG_PREFIX} INSERT BREAK: Heading 1 below {y_in_inches}\" at y={y/72:.2f}in."
                        )
                        if i > 0:
                            prev_para = doc.Paragraphs(i)
                            insert_rng = prev_para.Range.Duplicate
                            insert_rng.Collapse(c.wdCollapseStart)
                        else:
                            insert_rng = para.Range.Duplicate
                            insert_rng.Collapse(c.wdCollapseStart)
                        insert_rng.Select()
                        word.Selection.InsertBreak(c.wdPageBreak)
                        inserted_break = True
                        break

                    if (
                        is_bold_start
                        and y is not None
                        and next_y is not None
                        and next_y < y
                        and y > 72
                    ):
                        print(
                            f"  > {LOG_PREFIX} INSERT BREAK: Bold at start, next para above (spanned page), "
                            f"y={y/72:.2f}in, next_y={next_y/72:.2f}in"
                        )
                        if i > 0:
                            prev_para = doc.Paragraphs(i)
                            insert_rng = prev_para.Range.Duplicate
                            insert_rng.Collapse(c.wdCollapseStart)
                        else:
                            insert_rng = para.Range.Duplicate
                            insert_rng.Collapse(c.wdCollapseStart)
                        insert_rng.Select()
                        word.Selection.InsertBreak(c.wdPageBreak)
                        inserted_break = True
                        break

                    print(f"  > {LOG_PREFIX} No action.")
                    i += 1

                doc.Save()
                doc.Close(False)
                del doc
                time.sleep(0.4)
            except com_error as e:
                print(f"{LOG_PREFIX} COM error encountered: {e}. Retrying after sleep...")
                time.sleep(0.75)
                pythoncom.CoInitialize()
                continue
            pass_count += 1

        word.Quit()
