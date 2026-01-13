import os
import re
import logging
import tempfile
from typing import Any, Callable, Optional
import pythoncom
import win32com.client as win32
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Mm
from PIL import Image
import copy
from jinja2 import Undefined

logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)
if not logger.hasHandlers():
    handler = logging.StreamHandler()
    formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    handler.setFormatter(formatter)
    logger.addHandler(handler)

REPEATABLE_FIELDS = [
    "systemDesc",
    "systemoption",
    "riskdescription",
    "costSheet",
    "customerSpecifications",
    # Add more as needed
]

def flatten_jinja_repeatables(data: dict, repeatable_fields: list) -> dict:
    """
    Flattens repeatable nested fields for Jinja2 templates.
    For each repeatable field (e.g., systemDesc),
    {'name': {'1': 'foo'}, 'description': {'1': 'bar'}}
    becomes:
    {'name.1': 'foo', 'description.1': 'bar'}
    This modifies the input dict in place and returns it for convenience.
    """
    logger.info("[flatten_jinja_repeatables] Flattening repeatable fields...")
    for field in repeatable_fields:
        if field in data and isinstance(data[field], dict):
            new_block = {}
            for subfield, subdict in data[field].items():
                if isinstance(subdict, dict):
                    for idx, value in subdict.items():
                        new_block[f"{subfield}.{idx}"] = value
                        logger.debug(f"[flatten_jinja_repeatables] Flattened: {field}.{subfield}.{idx} -> {value}")
                else:
                    new_block[subfield] = subdict
            data[field] = new_block
            logger.info(f"[flatten_jinja_repeatables] Field '{field}' flattened")
    return data

class WordSubmissionHandler:
    """
    Handles Word document submission based on a template.

    Steps:
        1. Load Word template from disk (both python-docx and DocxTemplate)
        2. Render with Jinja2 context via DocxTemplate
        3. Save rendered temp file, reload with python-docx
        4. Replace image markers
        5. (Optional) Insert page breaks via Word COM
        6. Save final docx
    """

    def __init__(self, template_path: str, status_callback: Optional[Callable[[str], None]] = None) -> None:
        self.template_path = template_path
        self.status_callback = status_callback
        self.doc: Optional[Document] = None
        self.tpl: Optional[DocxTemplate] = None
        self.context: dict = {}
        self._send_status(f"WordSubmissionHandler initialized with template: {template_path}")

    def _send_status(self, message: str) -> None:
        if self.status_callback:
            self.status_callback(message)
        logger.info(message)

    @staticmethod
    def _unflatten_dict(flat: dict) -> dict:
        """
        Converts a flat dict with dot.notation keys to a nested dict.
        """
        result = {}
        for key, value in flat.items():
            parts = key.split('.')
            d = result
            for i, part in enumerate(parts):
                if i < len(parts) - 1:
                    if part not in d or not isinstance(d[part], dict):
                        d[part] = {}
                    d = d[part]
                else:
                    d[part] = value
        # If everything is under a "data" key, output with that root.
        if "data" in result and isinstance(result["data"], dict):
            return {"data": result["data"]}
        else:
            return result

    @staticmethod
    def _flatten_dict(d: dict, parent_key: str = '', sep: str = '.') -> dict:
        """
        Flattens a nested dict to a dict with dot.notation keys.
        """
        items = []
        for k, v in d.items():
            new_key = f"{parent_key}{sep}{k}" if parent_key else k
            if isinstance(v, dict):
                items.extend(WordSubmissionHandler._flatten_dict(v, new_key, sep=sep).items())
            else:
                items.append((new_key, v))
        return dict(items)

    def load_template(self) -> None:
        """
        Load both python-docx Document and docxtpl.DocxTemplate from the template path.
        """
        self._send_status("Loading Word template...")
        if not os.path.isfile(self.template_path):
            raise FileNotFoundError(f"Template file not found: {self.template_path}")
        try:
            self.doc = Document(self.template_path)
            self.tpl = DocxTemplate(self.template_path)
            self._send_status("Template loaded successfully.")
        except Exception as e:
            logger.exception("Error loading template.")
            raise e

    def analyze_template_and_context(self, context: dict):
        """
        Analyze the template to determine which variables are expected vs. provided.
        """
        self._send_status("==== [DEBUG] Analyzing template/context ====")
        tpl = self.tpl
        template_vars = tpl.get_undeclared_template_variables()
        self._send_status(f"Template expects variables: {sorted(template_vars)}")

        flat_context = self._flatten_dict(context)
        context_keys = set(flat_context.keys())
        self._send_status(f"Context provides keys: {sorted(context_keys)}")

        # Compare template_vars to context_keys (deep match)
        missing_in_context = set()
        used_and_present = set()
        for var in template_vars:
            if any(key == var or key.startswith(f"{var}.") or key.startswith(f"{var}[") for key in context_keys):
                used_and_present.add(var)
            else:
                missing_in_context.add(var)

        unused_in_template = set()
        for key in context_keys:
            key_root = key.split('.')[0]
            if key_root not in template_vars:
                unused_in_template.add(key)

        self._send_status(f"Missing in context: {sorted(missing_in_context)}")
        self._send_status(f"Unused in template: {sorted(unused_in_template)}")
        self._send_status(f"Used and present: {sorted(used_and_present)}")

        print("\n=== TEMPLATE/CONTEXT ANALYSIS ===")
        print("Template expects variables:", sorted(template_vars))
        print("Context provides keys:", sorted(context_keys))
        print("Missing in context:", sorted(missing_in_context))
        print("Unused in template:", sorted(unused_in_template))
        print("Used and present:", sorted(used_and_present))
        print("====================\n")

        return {
            'missing_in_context': missing_in_context,
            'unused_in_template': unused_in_template,
            'used_and_present': used_and_present,
        }

    def render_template(self, context: dict) -> None:
        """
        Render the template with context using docxtpl's default Jinja2 environment.
        Extremely verbose debug: logs context structure, types, template file checks, exceptions, and output .docx bytes.
        Also analyzes template/context variable matching.
        """
        import pprint
        import traceback
        import zipfile

        self._send_status("==== [DEBUG] Starting render_template ====")

        # 1. Log template file path and initial bytes
        self._send_status(f"Template path: {self.template_path}")
        try:
            with open(self.template_path, "rb") as f:
                bytes_head = f.read(200)
                self._send_status(f"First 200 bytes of template file: {bytes_head!r}")
        except Exception as e:
            self._send_status(f"ERROR reading template file: {e}")

        # 2. Log context structure and type of every value (deep)
        self._send_status("==== [DEBUG] Context Structure ====")
        self._send_status(pprint.pformat(context, indent=2, width=120))

        def check_types(obj, prefix=""):
            if isinstance(obj, dict):
                for k, v in obj.items():
                    check_types(v, f"{prefix}.{k}" if prefix else k)
            elif isinstance(obj, list):
                for i, v in enumerate(obj):
                    check_types(v, f"{prefix}[{i}]")
            else:
                self._send_status(f"Type at {prefix}: {type(obj).__name__} | Value: {repr(obj)[:120]}")
        check_types(context)

        # 3. OEE formatting as before
        for key, value in context.items():
            if "oee" in key.lower() and isinstance(value, (int, float)):
                context[key] = f"{value}%"

        # --- ANALYZE TEMPLATE/CONTEXT VARIABLES ---
        try:
            if self.tpl is None:
                raise ValueError("Template not loaded. Call load_template() first.")
            self.analyze_template_and_context(context)
        except Exception as e:
            self._send_status(f"ERROR during template/context analysis: {e}")

        # 4. Flatten repeatable fields just before rendering
        try:
            self._send_status("==== [DEBUG] Calling flatten_jinja_repeatables ====")
            if 'data' in context:
                flatten_jinja_repeatables(context['data'], REPEATABLE_FIELDS)
            else:
                flatten_jinja_repeatables(context, REPEATABLE_FIELDS)

            self._send_status("==== [DEBUG] Calling tpl.render(context) ====")
            self.tpl.render(context)
            self._send_status("tpl.render completed successfully.")
        except Exception as e:
            self._send_status("==== [DEBUG] Exception during tpl.render! ====")
            self._send_status(str(e))
            self._send_status(traceback.format_exc())
            raise

        # 5. Save to temp file and log file info
        try:
            temp_file = os.path.join(tempfile.gettempdir(), "temp_rendered.docx")
            self.tpl.save(temp_file)
            self._send_status(f"==== [DEBUG] tpl.save() complete. Rendered .docx at: {temp_file}")
            file_size = os.path.getsize(temp_file)
            self._send_status(f"Rendered .docx size: {file_size} bytes")
            with open(temp_file, "rb") as f:
                file_head = f.read(200)
                self._send_status(f"First 200 bytes of rendered .docx: {file_head!r}")

            # 6. List zip (docx) contents and head of document.xml
            with zipfile.ZipFile(temp_file, 'r') as z:
                file_list = z.namelist()
                self._send_status(f"Files in rendered .docx: {file_list}")
                if "word/document.xml" in file_list:
                    xml_bytes = z.read("word/document.xml")[:500]
                    self._send_status("First 500 bytes of word/document.xml:\n" +
                                      xml_bytes.decode("utf-8", errors="replace"))
                else:
                    self._send_status("word/document.xml not found in docx!")

            # 7. Reload with python-docx as before
            self.doc = Document(temp_file)
            self._send_status("==== [DEBUG] Finished render_template, temp_rendered.docx loaded successfully ====")

        except Exception as e:
            self._send_status("==== [DEBUG] Exception after tpl.save or while reading docx! ====")
            self._send_status(str(e))
            self._send_status(traceback.format_exc())
            raise

    def replace_image_markers(self, context: dict, image_width_mm: int = 50) -> None:
        """
        Replace image markers of form [[ Image: key ]] with actual images.
        """
        self._send_status("Replacing image markers...")
        image_marker_pattern = re.compile(r"\[\[\s*[Ii]mage:?\s*([^\]]+)\s*\]\]")
        try:
            if self.doc is None:
                raise ValueError("Document not loaded. Call load_template() first.")
            for paragraph in self.doc.paragraphs:
                matches = image_marker_pattern.findall(paragraph.text)
                for match in matches:
                    image_key = match.strip()
                    image_path = self._get_nested_context_value(context, image_key)
                    if not image_path:
                        logger.error(f"No image path found for key: {image_key}")
                        continue
                    if not os.path.isfile(image_path):
                        logger.error(f"Image file not found: {image_path}")
                        continue
                    # Remove marker text and insert picture
                    marker_text = f"[[ Image: {image_key} ]]"
                    paragraph.text = paragraph.text.replace(marker_text, "")
                    run = paragraph.add_run()
                    run.add_picture(image_path, width=Mm(image_width_mm))
                    logger.info(f"Inserted image for key '{image_key}' from path: {image_path}")
            self._send_status("Image markers replaced successfully.")
        except Exception as e:
            logger.exception("Error replacing image markers.")
            raise e

    def _get_nested_context_value(self, context: dict, key: str) -> Any:
        """
        Support dotted key notation for nested dicts: 'foo.bar.baz' -> context['foo']['bar']['baz']
        """
        parts = key.split('.')
        value = context
        for part in parts:
            if isinstance(value, dict) and part in value:
                value = value[part]
            else:
                return None
        return value

    def insert_page_breaks(self, file_path: str) -> None:
        """
        Insert page breaks before headings via Word COM automation.
        """
        self._send_status("Inserting page breaks...")
        try:
            pythoncom.CoInitialize()
            word = win32.gencache.EnsureDispatch('Word.Application')
            word.Visible = False
            doc = word.Documents.Open(file_path)
            page_setup = doc.PageSetup
            page_height = page_setup.PageHeight
            bottom_threshold = page_height * 0.75  # Bottom 25% threshold

            # For demonstration, iterate over paragraphs and insert a page break before each heading.
            for paragraph in doc.Paragraphs:
                style_name = paragraph.Style.NameLocal
                if "Heading" in style_name:
                    paragraph.Range.InsertBreak(Type=win32.constants.wdPageBreak)
            doc.Save()
            doc.Close()
            word.Quit()
            self._send_status("Page breaks inserted successfully.")
        except Exception as e:
            logger.exception("Error inserting page breaks.")
            raise e

    def save_document(self, output_path: str) -> None:
        """
        Save the final document to disk.
        """
        self._send_status(f"Saving document to: {output_path}")
        try:
            if self.doc is None:
                raise ValueError("No document loaded to save.")
            self.doc.save(output_path)
            self._send_status(f"Document saved successfully: {output_path}")
        except Exception as e:
            logger.exception("Error saving document.")
            raise e
